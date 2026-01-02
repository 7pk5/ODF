"""
File Indexer
Scans folders and extracts content from PDF, DOCX, and TXT files using parallel processing.
"""

import os
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document
import re
from tqdm import tqdm
import logging

# Suppress annoying PDFMiner warnings
logging.getLogger("pdfminer").setLevel(logging.ERROR)

class FileIndexer:
    def __init__(self):
        """Initialize the file indexer."""
        self.supported_extensions = {'.pdf', '.docx'}
        self.system_folders = {
            'C:\\WINDOWS',
            'C:\\PROGRAM FILES',
            'C:\\PROGRAM FILES (X86)',
            'C:\\SYSTEM32',
            'C:\\PROGRAMDATA',
            'C:\\USERS\\DEFAULT',
            'C:\\BOOT',
            'C:\\RECOVERY'
        }
    
    def scan_directory(self, folder_path):
        """
        Scans the folder and returns a list of supported files.
        Useful for getting a total count before processing.
        """
        if not os.path.exists(folder_path):
            raise ValueError(f"Folder does not exist: {folder_path}")
            
        all_files = []
        for root, dirs, files in os.walk(folder_path):
            # Filter directories in-place
            dirs[:] = [d for d in dirs if not self._should_skip_folder(os.path.join(root, d))]
            
            for file in files:
                _, ext = os.path.splitext(file.lower())
                if ext in self.supported_extensions:
                    all_files.append(os.path.join(root, file))
                    
        print(f"Found {len(all_files)} supported files to scan.")
        return all_files

    def process_files(self, file_paths, existing_ids=None):
        """
        Process a list of files in parallel (multiprocessing) and yield documents.
        """
        if existing_ids is None:
            existing_ids = set()
            
        print(f"Processing {len(file_paths)} files with {len(existing_ids)} existing IDs...")
        
        # Prepare arguments for map (file_path, existing_ids)
        # But we can't pass the huge set to every worker easily if not careful.
        # Actually, we can check the ID *before* submitting to executor if we can calculate ID cheaply.
        # Computing ID (md5 of path+mtime) is cheap. Extracting text is expensive.
        
        files_to_process = []
        skipped_count = 0
        
        for f in file_paths:
            # Pre-calculate ID to check if we can skip
            try:
                file_stats = os.stat(f)
                doc_id = hashlib.md5(f"{f}_{file_stats.st_mtime}".encode()).hexdigest()
                
                if doc_id in existing_ids:
                    skipped_count += 1
                    continue
                
                files_to_process.append(f)
            except:
                continue
                
        print(f"Skipping {skipped_count} already indexed files. Processing {len(files_to_process)} new/modified files.")

        print(f"Skipping {skipped_count} already indexed files. Processing {len(files_to_process)} new/modified files.")

        # Revert to ThreadPoolExecutor for Windows stability (ProcessPool requires strict entry point guards)
        with ThreadPoolExecutor(max_workers=min(32, os.cpu_count() * 4)) as executor:
            # Submit all tasks
            future_to_file = {executor.submit(self._process_single_path_independent, f): f for f in files_to_process}
            
            # Yield results as they complete
            for future in as_completed(future_to_file):
                try:
                    result = future.result()
                    if result:
                        yield result
                except Exception as e:
                    print(f"Error processing file: {e}")
        


    def _process_single_path_independent(self, file_path):
        """
        Process a single file: extract text and metadata.
        Designed to be called from a fresh instance in a worker process.
        """
        try:
            _, ext = os.path.splitext(file_path.lower())
            
            # We need to access extraction methods. Since we are in a fresh instance (wrapper created one),
            # we can use self._extract_content.
            content = self._extract_content(file_path, ext)
            
            if not content or not content.strip():
                return None
                
            file_stats = os.stat(file_path)
            doc_id = hashlib.md5(f"{file_path}_{file_stats.st_mtime}".encode()).hexdigest()
            
            return {
                "id": doc_id,
                "content": content,
                "metadata": {
                    "source": file_path,
                    "filename": os.path.basename(file_path),
                    "modified": file_stats.st_mtime,
                    "size": file_stats.st_size,
                    "type": ext
                }
            }
        except Exception as e:
            # print(f"Error in processing {file_path}: {e}")
            return None

    def _extract_content(self, file_path, extension):
        """Extract text content from a file based on its extension."""
        if extension == '.txt':
            return self._extract_txt_content(file_path)
        elif extension == '.pdf':
            return self._extract_pdf_content(file_path)
        elif extension == '.docx':
            return self._extract_docx_content(file_path)
        return ""
    
    def _extract_txt_content(self, file_path):
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return self._clean_text(file.read())
            except:
                continue
        return ""
    
    def _extract_pdf_content(self, file_path):
        try:
            content = extract_pdf_text(file_path)
            return self._clean_text(content)
        except:
            return ""
    
    def _extract_docx_content(self, file_path):
        try:
            doc = Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            return self._clean_text('\n'.join(full_text))
        except:
            return ""

    def _clean_text(self, text):
        if not text:
            return ""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove null bytes etc
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        return text.strip()[:100000] # Cap text size to avoid huge tokens

    def _should_skip_folder(self, folder_path):
        """Check if a folder should be skipped (system, hidden, dev envs)."""
        if not folder_path: return False
        
        name = os.path.basename(folder_path)
        
        # 1. Skip hidden directories
        if name.startswith('.'):
            return True
            
        # 2. Skip specific ignore patterns
        ignore_names = {
            'node_modules', 'site-packages', 'dist-info', '__pycache__', 
            'venv', 'env', 'odf_env', 'libs', 'include', 'scripts', 'bin', 'obj'
        }
        if name.lower() in ignore_names:
            return True
            
        # 3. Skip system folders (check start of path)
        norm_path = os.path.normpath(folder_path.upper())
        for sys in self.system_folders:
            if norm_path.startswith(sys):
                return True
                
        return False
